from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image

from docx import Document
from docx.shared import Pt, Inches

from datetime import datetime
from io import BytesIO

FILENAME_SEPARATOR = '_'
DATE_FORMAT = '%Y%m%d'
FOOTER_TEXT = "Generated on {date} by DISINFOX 🦊"
WEBSITE_TEXT = "https://disinfox.um.es"
INTERSECTION_VERTICAL_SPACING = 5






def export_incident_to_pdf(incident_stix_bundle):
    """
    Export an STIX2 disinformation incident to a PDF.
    """
    # Buffer para almacenar el PDF
    io = BytesIO()
    
    # Crear un documento con márgenes establecidos
    doc = SimpleDocTemplate(
        io, pagesize=letter,
        rightMargin=30, leftMargin=30,
        topMargin=30, bottomMargin=30
    )
    
    # Estilos predefinidos para texto
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    normal_style = styles['Normal']
    normal_style.leftIndent = 0  # Remove left indentation
    normal_style.firstLineIndent = 0  # Remove first-line indentation
    normal_style.spaceBefore = 0  # Remove space before the paragraph
    normal_style.spaceAfter = 0  # Remove space after the paragraph
    section_title_style = styles['Heading2']
    
    # Contenedor para los elementos del PDF
    elements = []
    
    # Title and logo, logo on the right side of the title
    title_logo_data = [
        [Paragraph("Incident Report", title_style), Image("static/images/logo.png", height=50, width=100)]
    ]
    # Table with page-wide length for the title and the logo
    title_logo_table = Table(title_logo_data)
    title_logo_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Alinear verticalmente
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),       # Logo a la derecha
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),       # Título a la izquierda,
        ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),  # Fuente en negrita para el título
        # No border and no paddings
        ('BOX', (0, 0), (-1, -1), 0, colors.white),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),

    ]))
    elements.append(title_logo_table)
    elements.append(Spacer(0, 24))

    
    # Procesar los datos de los incidentes
    incident = {}
    threat_actors = []
    locations = []
    attack_patterns = []
    for stix_object in incident_stix_bundle.get('objects', []):
        if stix_object['type'] == 'intrusion-set':
            incident = stix_object
        elif stix_object['type'] == 'threat-actor':
            threat_actors.append(stix_object.get('name', ''))
        elif stix_object['type'] == 'location':
            locations.append(stix_object.get('name', ''))
        elif stix_object['type'] == 'attack-pattern':
            attack_patterns.append(stix_object.get('name', ''))
    
    # Tabla de Intrusion Set
    if incident:
        data = [
            ['Name', Paragraph(incident.get('name', ''), normal_style)],
            ['Date', Paragraph(incident.get('first_seen', ''), normal_style)],
        ]
        # table that wraps the data
        table = Table(data, colWidths=[100, doc.width-100])
        table.setStyle(TableStyle([
            # Fox-like color
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#FFA500')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
            ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ]))
        elements.append(table)
        elements.append(Spacer(0, 24))
    # Descripción del incidente
    elements.append(Paragraph("Description:", section_title_style))
    elements.append(Paragraph(incident.get('description', ''), normal_style))
    elements.append(Spacer(0, INTERSECTION_VERTICAL_SPACING))
    
    # Lista de Locations
    if locations:
        elements.append(Paragraph("Locations:", section_title_style))
        elements.append(Paragraph(", ".join(locations), normal_style))
        elements.append(Spacer(0, INTERSECTION_VERTICAL_SPACING))
    
    # Lista de Threat Actors
    if threat_actors:
        elements.append(Paragraph("Threat Actors:", section_title_style))
        elements.append(Paragraph(", ".join(threat_actors), normal_style))
        elements.append(Spacer(0, INTERSECTION_VERTICAL_SPACING))
    
    # Lista de Attack Patterns
    if attack_patterns:
        elements.append(Paragraph("Attack Patterns:", section_title_style))
        elements.append(Paragraph(", ".join(attack_patterns), normal_style))
        elements.append(Spacer(0, INTERSECTION_VERTICAL_SPACING))
    
    def draw_footer(canvas, doc):
        """
        Function to draw the footer in the PDF
        """
        canvas.saveState()
        footer_text = FOOTER_TEXT.format(date=datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        website_text = WEBSITE_TEXT
        canvas.setFont('Helvetica', 10)
        canvas.drawString(30, 30, footer_text)
        canvas.drawString(30, 20, website_text)
        canvas.restoreState()
    # Construir el PDF# Construir el PDF con el footer
    doc.build(elements, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return io.getvalue()



def export_incident_to_word(incident_stix_bundle):
    """
    Export an STIX2 disinformation incident to a Word document.
    """
    # Create a new Word document
    doc = Document()

    # Title and logo (header)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    row = table.rows[0]
    title_cell = row.cells[0]
    logo_cell = row.cells[1]

    # Add title
    title_cell.text = "Incident Report"
    title_cell.paragraphs[0].style.font.size = Pt(16)
    title_cell.paragraphs[0].style.font.bold = True

    # Add logo (optional)
    try:
        logo_cell.add_paragraph().add_run().add_picture("static/images/logo.png", width=Inches(1.5))
    except FileNotFoundError:
        logo_cell.text = "Logo Missing"

    # Set alignment
    title_cell.paragraphs[0].alignment = 0  # Left align
    logo_cell.paragraphs[0].alignment = 2  # Right align

    # Process STIX bundle
    incident = {}
    threat_actors = []
    locations = []
    attack_patterns = []

    for stix_object in incident_stix_bundle.get('objects', []):
        if stix_object['type'] == 'intrusion-set':
            incident = stix_object
        elif stix_object['type'] == 'threat-actor':
            threat_actors.append(stix_object.get('name', ''))
        elif stix_object['type'] == 'location':
            locations.append(stix_object.get('name', ''))
        elif stix_object['type'] == 'attack-pattern':
            attack_patterns.append(stix_object.get('name', ''))

    # Add Incident Details Table
    if incident:
        doc.add_heading("Incident Details", level=2)
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = incident.get('name', '')
        table.cell(1, 0).text = "Date"
        table.cell(1, 1).text = incident.get('first_seen', '')

    # Add Description Section
    doc.add_heading("Description:", level=2)
    doc.add_paragraph(incident.get('description', ''))

    # Add Locations
    if locations:
        doc.add_heading("Locations:", level=2)
        doc.add_paragraph(", ".join(locations))

    # Add Threat Actors
    if threat_actors:
        doc.add_heading("Threat Actors:", level=2)
        doc.add_paragraph(", ".join(threat_actors))

    # Add Attack Patterns
    if attack_patterns:
        doc.add_heading("Attack Patterns:", level=2)
        doc.add_paragraph(", ".join(attack_patterns))

    # Draw Footer
    section = doc.sections[-1]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_text = footer_text = FOOTER_TEXT.format(date=datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    footer_paragraph.text = f"{footer_text}\n{WEBSITE_TEXT}"
    footer_paragraph.style.font.size = Pt(10)

    # Save Document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Posicionar el puntero al inicio del buffer
    return buffer.getvalue()